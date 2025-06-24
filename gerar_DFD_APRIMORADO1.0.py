import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime


MODELO_OLLAMA = "llama3:8b"


# =========================
# 🔹 Funções de apoio 🔹
# =========================

def gerar_descricao_ollama(prompt):
    cmd = ['ollama', 'run', MODELO_OLLAMA]
    try:
        process = subprocess.Popen(
            cmd,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8',
            errors='replace',
            text=True
        )
        stdout, stderr = process.communicate(input=prompt, timeout=300)
    except subprocess.TimeoutExpired:
        process.kill()
        print("⛔ Timeout: o processo demorou demais para responder.")
        return None

    if process.returncode != 0:
        print("⛔ Erro ao executar Ollama:", stderr.strip())
        return None
    return stdout.strip()


def configurar_estilo(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)


def adicionar_cabecalho(doc):
    section = doc.sections[0]
    header = section.header
    par = header.paragraphs[0]
    par.text = "COORDENADORIA DE TECNOLOGIA DA INFORMAÇÃO"
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = par.runs[0]
    run.font.size = Pt(9)
    run.bold = True
    run.font.name = 'Calibri'


def adicionar_rodape(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = (
        "Av. Dr. Hélio Ribeiro, nº 1000, Res. Paiaguás – Cuiabá-MT – CEP 78048-910 | "
        "Fone: (65) 3615-4811 | cti@detran.mt.gov.br | www.detran.mt.gov.br"
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(8)
    run.font.name = 'Calibri'


def adicionar_conceitos_gerais(doc):
    titulo = doc.add_paragraph("CONCEITOS GERAIS")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    texto = (
        "O Decreto Estadual nº 1.525/2022 em seu Art. 42 disciplinou que: O termo de referência é o "
        "documento elaborado a partir dos estudos técnicos preliminares, se houver, devendo conter os elementos "
        "necessários e suficientes, com nível de precisão adequado, para caracterizar o objeto da licitação, e ainda: "
        "I - definição do objeto, incluídos sua natureza, os quantitativos, o prazo do contrato e, se for o caso, a "
        "possibilidade de sua prorrogação; II - fundamentação da contratação, que consiste na referência aos estudos "
        "técnicos preliminares correspondentes ou, quando não for possível divulgar esses estudos, no extrato das "
        "partes que não contiverem informações sigilosas; III - descrição da solução como um todo, considerando "
        "todo o ciclo de vida do objeto; IV - requisitos da contratação; V - modelo de execução do objeto, que consiste "
        "na definição de como o contrato deverá produzir os resultados pretendidos desde o seu início até o seu "
        "encerramento; VI - modelo de gestão do contrato, que descreve como a execução do objeto será "
        "acompanhada e fiscalizada pelo órgão ou entidade; VII - critérios de medição e de pagamento; VIII - forma e "
        "critérios de seleção do contratado; IX - estimativas do valor da contratação, acompanhadas dos preços "
        "unitários referenciais, das memórias de cálculo e dos documentos que lhe dão suporte, com os parâmetros "
        "utilizados para a obtenção dos preços e para os respectivos cálculos, que devem constar de documento "
        "separado e classificado; X - adequação orçamentária; XI - indicação dos locais de execução dos serviços e das "
        "regras para recebimento provisório e definitivo, quando for o caso; XII - especificação da garantia exigida e "
        "das condições de manutenção e assistência técnica, quando for o caso; XIII - formas, condições e prazos de "
        "pagamento, bem como o critério de reajuste; XIV - principais obrigações do contratado e do contratante, "
        "inclusive com a eventual previsão da execução de logística reversa pelo contratado, se for o caso; e XV - "
        "sanções por descumprimentos das obrigações pactuadas, inclusive as obrigações prévias ao contrato.\n\n"
        "Em seu §1º ainda disciplina que para a definição do objeto, deverá ser utilizada a especificação do "
        "produto ou serviço existente no catálogo de especificações do Sistema de Aquisições Governamentais ou "
        "solicitada a sua inclusão quando se tratar de novos produtos ou serviços, observados os requisitos de "
        "qualidade, rendimento, compatibilidade, durabilidade e segurança. E por fim o §2º diz que o termo de "
        "referência deverá ser elaborado por servidor da área técnica, auxiliado pela área de contratação nos aspectos "
        "técnicos de compras públicas.\n\n"
        "Já no seu Art. 66 diz que os processos de aquisição de bens e de contratação de serviços e locação "
        "de bens móveis e imóveis serão autuados e instruídos em sua fase interna pelo menos com os seguintes "
        "documentos, dentre eles: I - documento de formalização de demanda com a justificativa para a "
        "contratação, termo de referência, projeto básico ou projeto executivo e, se for o caso, estudo técnico "
        "preliminar e análise de riscos.\n\n"
        "Desta forma, o presente documento visa ao levantamento de informações para elaboração das peças "
        "técnicas necessárias aos procedimentos de contratação pública."
    )
    par = doc.add_paragraph(texto)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def criar_tabela_cabecalho(doc):
    tabela = doc.add_table(rows=4, cols=4)
    tabela.style = 'Table Grid'

    tabela.cell(0, 0).text = "ÓRGÃO:"
    tabela.cell(0, 1).text = "DEPARTAMENTO ESTADUAL DE TRÂNSITO DE MATO GROSSO – DETRAN-MT"
    tabela.cell(0, 2).text = "MATRÍCULA:"
    tabela.cell(0, 3).text = "246679"

    tabela.cell(1, 0).text = "UNIDADE ORÇAMENTÁRIA:"
    tabela.cell(1, 1).text = "19301"
    tabela.cell(1, 2).text = "TELEFONE:"
    tabela.cell(1, 3).text = "(65) 3615-4811"

    tabela.cell(2, 0).text = "SETOR REQUISITANTE (UNIDADE/SETOR/DEPTO):"
    tabela.cell(2, 1).merge(tabela.cell(2, 3))
    tabela.cell(2, 1).text = "COORDENADORIA DE TECNOLOGIA DA INFORMAÇÃO"

    tabela.cell(3, 0).text = "RESPONSÁVEL PELA DEMANDA:"
    tabela.cell(3, 1).text = "DANILO VIEIRA DA CRUZ"
    tabela.cell(3, 2).text = "E-MAIL:"
    tabela.cell(3, 3).text = "danilocruz@detran.mt.gov.br"


# ===========================
# 🔥 Função principal
# ===========================

def gerar_dfd(descricao, tipo_objeto, forma_contratacao, estudo_tecnico, plano_contratacao):
    doc = Document()

    configurar_estilo(doc)
    adicionar_cabecalho(doc)
    adicionar_rodape(doc)
    adicionar_conceitos_gerais(doc)

    # ======== TÍTULO DO DOCUMENTO ========
    titulo = doc.add_paragraph("DOCUMENTO DE FORMALIZAÇÃO DA DEMANDA (DFD)")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    criar_tabela_cabecalho(doc)
    doc.add_paragraph()

    # ======= 1 - OBJETO =======
    doc.add_paragraph("1 - OBJETO (SOLUÇÃO PRELIMINAR):").runs[0].bold = True

    opcoes_objeto = {
        "Material de consumo": "(X)" if tipo_objeto == "consumo" else "( )",
        "Material permanente": "(X)" if tipo_objeto == "permanente" else "( )",
        "Equipamento de TI": "(X)" if tipo_objeto == "ti" else "( )",
        "Serviço não continuado": "(X)" if tipo_objeto == "servico_nao" else "( )",
        "Serviço continuado sem dedicação exclusiva de mão de obra": "(X)" if tipo_objeto == "servico_sem_exclusiva" else "( )",
        "Serviço continuado com dedicação exclusiva de mão de obra": "(X)" if tipo_objeto == "servico_com_exclusiva" else "( )",
    }
    for texto, marcado in opcoes_objeto.items():
        doc.add_paragraph(f" {marcado} {texto}")

    doc.add_paragraph()

    # ======= 2 - DESCRIÇÃO =======
    doc.add_paragraph("2 - DESCRIÇÃO SUCINTA DO OBJETO:").runs[0].bold = True
    doc.add_paragraph(descricao)
    doc.add_paragraph()

    # ======= 3 - FORMA DE CONTRATAÇÃO =======
    doc.add_paragraph("3 - FORMA DE CONTRATAÇÃO SUGERIDA:").runs[0].bold = True

    formas = {
        "Modalidades da Lei nº 14.133/21": "(X)" if forma_contratacao == "lei14133" else "( )",
        "Utilização à ARP - Órgão Participante": "(X)" if forma_contratacao == "arp" else "( )",
        "Adesão à ARP de outro Órgão": "(X)" if forma_contratacao == "adesao" else "( )",
        "Dispensa/Inexigibilidade": "(X)" if forma_contratacao == "dispensa" else "( )",
    }

    for texto, marcado in formas.items():
        doc.add_paragraph(f" {marcado} {texto}")

    doc.add_paragraph()

    # ======= 4 - ESTUDO TÉCNICO =======
    doc.add_paragraph("4 - NECESSIDADE DE ESTUDO TÉCNICO PRELIMINAR E ANÁLISE DE RISCOS:").runs[0].bold = True

    doc.add_paragraph(f" (X) {estudo_tecnico}" if estudo_tecnico == "NÃO" else f" (X) SIM")
    doc.add_paragraph("Justificativa: Não se aplica.")
    doc.add_paragraph()

    # ======= 5 - PLANO DE CONTRATAÇÕES =======
    doc.add_paragraph("5 - OS OBJETOS A SEREM ADQUIRIDOS/CONTRATADOS ESTÃO PREVISTOS NO PLANO DE CONTRATAÇÕES ANUAL:").runs[0].bold = True

    doc.add_paragraph(f" (X) {plano_contratacao}" if plano_contratacao == "SIM" else f" (X) NÃO")
    doc.add_paragraph("Justificativa: Não se aplica.")

    doc.add_paragraph()

    # ======= DATA E ASSINATURA =======
    data = datetime.now().strftime("Cuiabá-MT, %d de %B de %Y")
    data_par = doc.add_paragraph(data)
    data_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph("\n\n")
    assinatura = doc.add_paragraph("DANILO VIEIRA DA CRUZ")
    assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cargo = doc.add_paragraph("RESPONSÁVEL PELA AÇÃO NO PTA")
    cargo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    nome_arquivo = f"DFD_DETRANMT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    print(f"✅ Documento '{nome_arquivo}' criado com sucesso!")


# ===========================
# 🔥 FLUXO INTERATIVO
# ===========================

def fluxo_interativo():
    print("Informe os dados para gerar o DFD.\n")
    item = input("Item: ").strip()
    quantidade = input("Quantidade: ").strip()
    finalidade = input("Finalidade ou descrição resumida: ").strip()

    print("\nEscolha o tipo de objeto:")
    print("[1] Material de consumo")
    print("[2] Material permanente")
    print("[3] Equipamento de TI")
    print("[4] Serviço não continuado")
    print("[5] Serviço sem dedicação exclusiva de mão de obra")
    print("[6] Serviço com dedicação exclusiva de mão de obra")
    tipo = input("Digite o número: ").strip()
    tipos = {
        "1": "consumo",
        "2": "permanente",
        "3": "ti",
        "4": "servico_nao",
        "5": "servico_sem_exclusiva",
        "6": "servico_com_exclusiva"
    }
    tipo_objeto = tipos.get(tipo, "consumo")

    print("\nEscolha a forma de contratação:")
    print("[1] Modalidades da Lei nº 14.133/21")
    print("[2] Utilização à ARP - Órgão Participante")
    print("[3] Adesão à ARP de outro Órgão")
    print("[4] Dispensa/Inexigibilidade")
    forma = input("Digite o número: ").strip()
    formas = {
        "1": "lei14133",
        "2": "arp",
        "3": "adesao",
        "4": "dispensa"
    }
    forma_contratacao = formas.get(forma, "lei14133")

    estudo_tecnico = "NÃO"
    plano_contratacao = "SIM"

    # ==== Prompt para Ollama ====
    prompt = f"""
Você é um assistente que gera textos formais para documentos oficiais de requisição de materiais do DETRAN-MT.

Por favor, escreva uma descrição clara, formal e objetiva para justificar a aquisição de {quantidade} unidades do item '{item}', 
que será utilizado para {finalidade}.

O texto deve seguir o padrão institucional, com linguagem técnica, direta e adequada para documentos públicos de aquisição.
"""

    print("\n🧠 Gerando descrição com o modelo...\n")
    descricao_gerada = gerar_descricao_ollama(prompt)

    if descricao_gerada:
        gerar_dfd(descricao_gerada, tipo_objeto, forma_contratacao, estudo_tecnico, plano_contratacao)
    else:
        print("❌ Não foi possível gerar a descrição com Ollama.")


if __name__ == "__main__":
    fluxo_interativo()
