import subprocess
from docx import Document
from datetime import datetime

def gerar_descricao_ollama(prompt):
    cmd = ['ollama', 'run', 'llama3.2:1b']
    try:
        process = subprocess.Popen(
            cmd,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8',
            errors='replace',
            text=True
        )
        stdout, stderr = process.communicate(input=prompt, timeout=120)
    except subprocess.TimeoutExpired:
        process.kill()
        print("Timeout: o processo demorou demais para responder.")
        return None

    if process.returncode != 0:
        print("Erro ao executar Ollama:", stderr.strip())
        return None
    return stdout.strip()

def criar_docx_com_cabecalho(nome_arquivo, descricao):
    doc = Document()

    # Cabeçalho fixo e institucional
    doc.add_paragraph("DOCUMENTO DE FORMALIZAÇÃO DA DEMANDA (DFD)")
    doc.add_paragraph("ÓRGÃO: DEPARTAMENTO ESTADUAL DE TRÂNSITO DE MATO GROSSO – DETRAN-MT")
    doc.add_paragraph("UNIDADE ORÇAMENTÁRIA: 19301")
    doc.add_paragraph("SETOR REQUISITANTE: COORDENADORIA DE TECNOLOGIA DA INFORMAÇÃO")
    doc.add_paragraph("RESPONSÁVEL PELA DEMANDA: DANILO VIEIRA DA CRUZ")
    doc.add_paragraph("MATRÍCULA: 246679")
    doc.add_paragraph("E-MAIL: danilocruz@detran.mt.gov.br")
    doc.add_paragraph("TELEFONE: (65) 3615-4811")
    doc.add_paragraph("\n")

    # Espaço para a descrição variável gerada pelo Ollama
    doc.add_paragraph("DESCRIÇÃO SUCINTA DO OBJETO:")
    doc.add_paragraph(descricao)

    doc.save(nome_arquivo)
    print(f"Documento '{nome_arquivo}' criado com sucesso!")

def fluxo_interativo():
    print("Informe os dados para gerar o documento de requisição.\n")
    item = input("Item: ").strip()
    quantidade = input("Quantidade: ").strip()
    finalidade = input("Finalidade ou descrição resumida: ").strip()

    prompt = f"""
Você é um assistente que gera textos formais para documentos oficiais de requisição de materiais do DETRAN-MT.

Por favor, escreva uma descrição clara, formal e objetiva para justificar a aquisição de {quantidade} unidades do item '{item}', 
que será utilizado para {finalidade}.

O texto deve seguir o padrão institucional, com linguagem técnica, direta e adequada para documentos públicos de aquisição.
"""

    print("\nGerando descrição com o modelo...\n")
    descricao_gerada = gerar_descricao_ollama(prompt)

    if descricao_gerada:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"Documento_Requisicao_DETRANMT_{timestamp}.docx"
        criar_docx_com_cabecalho(nome_arquivo, descricao_gerada)
    else:
        print("Não foi possível gerar a descrição com Ollama.")

if __name__ == "__main__":
    fluxo_interativo()
