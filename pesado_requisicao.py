import subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime


def gerar_descricao_ollama(prompt):
    cmd = ['ollama', 'run', 'llama3']
    try:
        process = subprocess.Popen(
            cmd,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8',
            errors='replace',
            text=True
        )
        stdout, stderr = process.communicate(input=prompt, timeout=300)
    except subprocess.TimeoutExpired:
        process.kill()
        print("Timeout: o processo demorou demais para responder.")
        return None

    if process.returncode != 0:
        print("Erro ao executar Ollama:", stderr.strip())
        return None
    return stdout.strip()


def criar_docx_formatado(nome_arquivo, descricao):
    doc = Document()

    # ===== TÍTULO CENTRAL =====
    titulo = doc.add_paragraph("DOCUMENTO DE FORMALIZAÇÃO DA DEMANDA (DFD)")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo_run = titulo.runs[0]
    titulo_run.font.size = Pt(16)
    titulo_run.bold = True

    doc.add_paragraph()  # Linha em branco

    # ===== DADOS INSTITUCIONAIS =====
    dados = {
        "ÓRGÃO": "DEPARTAMENTO ESTADUAL DE TRÂNSITO DE MATO GROSSO – DETRAN-MT",
        "UNIDADE ORÇAMENTÁRIA": "19301",
        "SETOR REQUISITANTE": "COORDENADORIA DE TECNOLOGIA DA INFORMAÇÃO",
        "RESPONSÁVEL PELA DEMANDA": "DANILO VIEIRA DA CRUZ",
        "MATRÍCULA": "246679",
        "E-MAIL": "danilocruz@detran.mt.gov.br",
        "TELEFONE": "(65) 3615-4811"
    }

    for chave, valor in dados.items():
        par = doc.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = par.add_run(f"{chave}: ")
        run.bold = True
        par.add_run(valor)

    doc.add_paragraph()  # Espaço

    # ===== SEÇÃO DE DESCRIÇÃO DO OBJETO =====
    secao = doc.add_paragraph()
    run_secao = secao.add_run("2 - DESCRIÇÃO SUCINTA DO OBJETO:\n")
    run_secao.bold = True
    run_secao.font.size = Pt(12)

    par_desc = doc.add_paragraph(descricao)
    par_desc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    par_desc.paragraph_format.space_after = Pt(10)

    # ===== DATA E ASSINATURA =====
    doc.add_paragraph()
    data_atual = datetime.now().strftime("Cuiabá-MT, %d de %B de %Y")
    p_data = doc.add_paragraph(data_atual)
    p_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph("\n\n")
    p_assinatura = doc.add_paragraph("DANILO VIEIRA DA CRUZ")
    p_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_cargo = doc.add_paragraph("RESPONSÁVEL PELA AÇÃO NO PTA")
    p_cargo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== SALVAR DOCUMENTO =====
    doc.save(nome_arquivo)
    print(f"Documento '{nome_arquivo}' criado com sucesso!")


def fluxo_interativo():
    print("Informe os dados para gerar o documento de requisição.\n")
    item = input("Item: ").strip()
    quantidade = input("Quantidade: ").strip()
    finalidade = input("Finalidade ou descrição resumida: ").strip()

    # ===== PROMPT PARA OLLAMA =====
    prompt = f"""
Você é um assistente que gera textos formais para documentos oficiais de requisição de materiais do DETRAN-MT.

Por favor, escreva uma descrição clara, formal e objetiva para justificar a aquisição de {quantidade} unidades do item '{item}', 
que será utilizado para {finalidade}.

O texto deve seguir o padrão institucional, com linguagem técnica, direta e adequada para documentos públicos de aquisição.
"""

    print("\n🧠 Gerando descrição com o modelo...\n")
    descricao_gerada = gerar_descricao_ollama(prompt)

    if descricao_gerada:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"Documento_Requisicao_DETRANMT_{timestamp}.docx"
        criar_docx_formatado(nome_arquivo, descricao_gerada)
    else:
        print("❌ Não foi possível gerar a descrição com Ollama.")


if __name__ == "__main__":
    fluxo_interativo()
